from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Your JSON data
# output_json = [
#     {"module_name": "Module 1", "Topic": "Topic 1", "sub-section": "Subsection 1", "course_content": "Content 1"},
#     {"module_name": "Module 2", "Topic": "Topic 2", "sub-section": "Subsection 2", "course_content": "Content 2"},
#     # Add more data as needed
# ]

#Create a new DOCX document
def create_doc_course(output_json,doc_name):
    doc = Document()

    # Iterate through the JSON data and add it to the document with headings
    for data in output_json:
        module_name = data["module_name"]
        sub_section = data["sub-section"]
        
        # Add the module name as a heading
        heading = doc.add_heading(module_name, level=1)
        heading = doc.add_heading(sub_section, level=3)
        
        # Add other data as paragraphs
        doc.add_paragraph(f"Course Content: {data['content']}")
        
        # Add an empty line between entries
        doc.add_paragraph("")
        doc.add_paragraph("")

    # Save the DOCX file
    doc.save(f"{doc_name}_document.docx")

def create_doc_exam(output,course_name):
    doc = Document()

    doc.add_heading(f"Exam {course_name}", level=1)
    doc.add_paragraph(f"{output}")

    
    # Add an empty line between entries
    doc.add_paragraph("")
    doc.add_paragraph("")
    #doc.add_heading("Solutions", level=3)


    # Save the DOCX file
    doc.save(f"{course_name}_exam_document.docx")