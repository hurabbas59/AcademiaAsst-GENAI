from exam_creator import ExamCreator
from course_generator import CourseGenerator
import asyncio
import json
from docx import Document
import os
from doc_creator import create_doc_exam, create_doc_course
from parsers import parse_json

obj_course_generator = CourseGenerator()
obj_exam_creator = ExamCreator()


def create_doc(data,folder):
    doc = Document()

    # Add Course Title
    doc.add_heading(f"Course Title: {data['Course title']}", level=1)

    # Add Course Code
    doc.add_heading(f"Course Code: {data['Course code']}", level=2)

    # Add Credits Hours
    doc.add_heading(f"Credit Hours: {data['Credit Hours']}", level=2)

    # Add Course Description
    doc.add_heading("Course Description", level=2)
    doc.add_paragraph(data["Course description"])

    # Add Learning Outcomes
    doc.add_heading("Learning Outcomes", level=2)
    counter = 1
    for outcome in data["Learning outcome"]:
        doc.add_paragraph(f"{counter}- {outcome}")
        counter+=1

    # Add Course Plan
    doc.add_heading("Course Plan", level=1)
    for module in data["Course Plan"]:
        doc.add_heading(module["Module_name"],level=3)
        doc.add_paragraph(f"Description: {module['Description']}")

    if "Lab Plan" in list(data.keys()):
        if data["Lab Plan"]:
            doc.add_heading("Lab Plan", level=1)
            for module in data["Lab Plan"]:
                doc.add_heading(module["Module_name"],level=3)
                doc.add_paragraph(f"Description: {module['Description']}")


    # Save the document to a file
    doc_name = f"{data['Course title']}.docx"
    
    doc.save(os.path.join(folder,doc_name))


# # here I am making changesss----------------------------------------------------  


async def generate_module(course_title: str,semester: int):
 
    modules = obj_course_generator.func_module_creator(course_title=course_title,semester=semester)
    modules_json_list = parse_json(modules)
    
    return modules_json_list

async def module_outline_future(module_name: str,module_description: str, semester: int):
    response = await asyncio.get_event_loop().run_in_executor(
        None,
        lambda: obj_course_generator.func_module_content_outline_creator(module_name=module_name,
                                                                         module_description=module_description,
                                                                         semester=semester)
    )
    
    return response
    
async def generate_module_outline(modules_info: list, semester: int):
    tasks = []

    for module_info in modules_info:
        task = asyncio.ensure_future(
            module_outline_future(module_name=module_info["module_name"],
                                  module_description=module_info["description"],
                                  semester=semester
                                 )
        )
        tasks.append(task)

    # Gather results in the same order as tasks were created
    results = await asyncio.gather(*tasks)

    # Parse and append results in the same order
    modules_outline = []
    for result in results:
        try:
            response = parse_json(result)
            modules_outline.append(response)
        except Exception:
            modules_outline.append(result)

    return modules_outline


# CHange function Module Creator
async def generate_module_future(course_name: str,course_code: str, credit_hours: str):
    response = await asyncio.get_event_loop().run_in_executor(
        None,
        lambda: obj_course_generator.func_module_outline_creator(course_name=course_name,course_code=course_code,credit_hours=credit_hours)
    )
    
    return response

async def generate_course_module(file_name):
    
    # Read file
    
    with open(file_name,'r') as file:
        data = json.load(file)
    # Create tasks
    
    tasks = []
    course_names = []
    course_codes = []
    credit_hours = []
    semesters = [] 
    for courses in data[:1]:
        semester = courses['semester']
        #print(semester)

        for sub in courses['courses']:
            course_titile = sub['course name']
            
            course_code = sub['course code']
            credit_hour =  sub["credit hrs"]            
            course_names.append(course_titile)
            course_codes.append(course_code)
            credit_hours.append(credit_hour)
            semesters.append(semester)
            
            tasks.append(
                asyncio.ensure_future(
                    generate_module_future(course_name=course_titile,course_code=course_code,credit_hours=credit_hour)
                )
            )
            
    
    # Gather results in the same order as tasks were created
    results = await asyncio.gather(*tasks)
    response = []
    semesters_lookup = {}
    sem_counter = 0
    for idx,result in enumerate(results):
        try:

            data = json.loads(result)
            response.append(data)
        except Exception:
            try:

                print("Exception Occured")
                data = json.loads(obj_course_generator.json_converter(result))
                response.append(data)
            except Exception:
                continue
        
        
        if semesters[idx] not in semesters_lookup:
            sem_counter+=1
            semesters_lookup[semesters[idx]] = True
            os.mkdir(f"{semester}")
            create_doc(data=data,folder=semester)
        else:
            create_doc(data=data,folder=f"{semester}")

    return response



# Exam--------
            

async def generat_exam(course_name,course_type,difficulty_level,exam_time,semester,topics_to_include):
    exam_response = obj_exam_creator.func_generate_exam(course_name,course_type,difficulty_level,exam_time,semester,topics_to_include)  
    print(exam_response)
    create_doc_exam(output=exam_response,course_name=course_name)
    
    return exam_response




# ------- ---------- Content Creator ---------------


async def func_theoretical_content_creator_future(sub_section,detailed_description,semester):
        response = await asyncio.get_event_loop().run_in_executor(
            None,
            lambda: obj_course_generator.func_theoretical_content_creator(sub_section,detailed_description,semester)
        )
        
        return response
    
async def func_technical_content_creator_future(sub_section,detailed_description,semester):
        response = await asyncio.get_event_loop().run_in_executor(
            None,
            lambda: obj_course_generator.func_technical_content_creator(sub_section,detailed_description,semester)
        )
        
        return response
    
async def generate_module_content(course_outline:list,course_title: str,modules_info: list, semester: int, course_type: str):
    
    module_names = []
    tasks = []
    results = []
    for _module in modules_info[:2]:
        module_names.append(_module['module_name'])

    counter = 0 
    for outline in course_outline[:2]:
        
        topic = outline[0]['Topic']
        topic_content = outline[0]["Topic Content"]
        
        for content_des in topic_content:
            sub_section = content_des['sub-section']
            sub_section_description = content_des['detailed_description']
            
            if course_type.lower()=='theoretical':
                task = asyncio.ensure_future(
                    func_theoretical_content_creator_future(sub_section=sub_section,
                                                            detailed_description=sub_section_description,
                                                            semester=semester)
                )
                results.append({
                    "Module":topic,
                    "sub-section":sub_section,
                    
                })
                
                tasks.append(task)
                
            if course_type.lower()=='technical':
                task = asyncio.ensure_future(
                    func_technical_content_creator_future(sub_section=sub_section,
                                                          detailed_description=sub_section_description,
                                                          semester=semester)
                )
                results.append({
                    "module_name":topic,
                    "sub-section":sub_section,
                    
                })
                tasks.append(task)
            
            
    completed_tasks, _ = await asyncio.wait(tasks)
    
    i = 0 
    for task in completed_tasks:
        
        results[i]['content'] = task.result()
        i+=1
    print(results)
    # Create doc
    create_doc_course(results,doc_name=course_title)
    return results